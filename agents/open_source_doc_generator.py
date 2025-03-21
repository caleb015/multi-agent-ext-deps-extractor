from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class OpenSourceDocGenerator:
    """
    Generates a DOCX document based on the extracted open-source dependencies.
    """
    
    def __init__(self, repo_path: str, researcher_output: str, app_name: str):
        """
        :param repo_path: Path to the repository where output files will be saved.
        :param researcher_output: Path to the JSON file containing researched dependency data.
        :param app_name: Name of the application.
        """
        self.repo_path = repo_path
        self.researcher_output = researcher_output
        self.app_name = app_name
        self.company_name = os.getenv("COMPANY_NAME", "Your Company")  # Load company name from .env
        self.company_email = os.getenv("COMPANY_EMAIL", "contact@company.com")  # Load company email from .env
        self.shed_dir = os.path.join(repo_path, ".shed")
        os.makedirs(self.shed_dir, exist_ok=True)
        self.doc_filename = os.path.join(self.shed_dir, "open_source_declaration.docx")
        
    def load_dependencies(self):
        """Loads open-source dependencies from the JSON output of the Web Researcher Agent."""
        with open(self.researcher_output, "r") as f:
            dependencies = json.load(f)
        
        return [dep for dep in dependencies if dep.get("is_open_source", False)]
    
    def generate_document(self):
        """Creates a DOCX document with open-source dependency details."""
        dependencies = self.load_dependencies()
        
        doc = Document()
        
        # Add a centered header with company and app name
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header.add_run(f"{self.company_name} - {self.app_name}")
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_heading("Open Source Declaration", level=1)
        
        # Add paragraphs with bold app name
        paragraph = doc.add_paragraph(
            "This document contains information about the open-source software components used in "
        )
        paragraph.add_run(self.app_name).bold = True
        paragraph.add_run(
            ". It includes details of the applicable licenses, acknowledgments required by the respective licensors, "
            "and information on obtaining the source code (where applicable). This list of open-source code has been "
            "compiled by reference to third-party software incorporated into the "
        )
        paragraph.add_run(self.app_name).bold = True
        paragraph.add_run(
            " as of the date the list was created. Therefore, this list may be updated from time to time.\n\n"
            "All information herein is provided \"as is\". "
        )
        paragraph.add_run(self.app_name).bold = True
        paragraph.add_run(
            " and its suppliers make no warranties, express or implied, regarding this list and its accuracy and completeness "
            "or the results that may be obtained from the use or distribution of the list. By using or distributing this list, "
            "you agree that in no event shall "
        )
        paragraph.add_run(self.app_name).bold = True
        paragraph.add_run(
            " be liable for any damages resulting from any use or distribution of this list, including, without this list being "
            "exclusive, special, consequential, incidental, or any other direct or indirect damages."
        )
        
        doc.add_heading("Open Source Components and Licenses", level=2)
        
        for idx, dep in enumerate(dependencies, start=1):
            doc.add_paragraph(
                f"{idx}. {dep['package_name']} (License: {dep['license']['name']})"
            )
            doc.add_paragraph(f"   - License URL: {dep['license'].get('url', 'N/A')}")
            doc.add_paragraph(f"   - Versions: {', '.join(dep['installed_versions'])}")
            
        doc.add_heading("How to Obtain Source Code", level=2)
        paragraph = doc.add_paragraph(
            "For components where the license requires providing source code, contact "
        )
        paragraph.add_run(self.company_name).bold = True
        paragraph.add_run(" at ")
        paragraph.add_run(self.company_email).bold = True
        paragraph.add_run(".")
        
        doc.add_heading("Acknowledgments", level=2)
        doc.add_paragraph("This product includes software developed by various open-source contributors.")
        
        doc.add_heading("Additional Information", level=2)
        paragraph = doc.add_paragraph()
        paragraph.add_run(self.app_name).bold = True
        paragraph.add_run(
            " is committed to supporting the open-source community and complying with all applicable open-source licenses. "
            "We are grateful to the developers and contributors of these open-source projects that have made this product possible."
        )
        paragraph = doc.add_paragraph(
            "For full source code and further information about the open-source components used in this product, please contact "
        )
        paragraph.add_run(self.company_name).bold = True
        paragraph.add_run(" at ")
        paragraph.add_run(self.company_email).bold = True
        paragraph.add_run(".")
        
        doc.save(self.doc_filename)
        print(f"✅ Document generated: {self.doc_filename}")
        
    def run(self):
        """Executes the document generation process."""
        self.generate_document()